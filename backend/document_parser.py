"""
backend/document_parser.py — Multi-Format Extraction Layer
===========================================================
Universal Document Ingestion (Phase 5).

Supports extracting text and structuring candidate data from live uploaded
resumes in PDF, DOCX, and TXT formats.

Phase 5b update
---------------
``structure_candidate_data`` now also extracts (best-effort, regex/keyword
heuristics — no external NLP dependency required):

* ``profile.github``    — GitHub profile URL, if present.
* ``profile.linkedin``  — LinkedIn profile URL, if present.
* ``education``         — list of detected degree / institution lines.
* ``projects``          — list of detected project entries (name + blurb).
* ``certifications``    — list of detected certification names.

All of these default to an empty list / empty string rather than raising,
so the downstream scoring & dashboard pipeline never breaks on a sparse
resume.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any, Union

logger = logging.getLogger(__name__)


class UniversalParser:
    """Universal parser for resume documents (PDF, DOCX, TXT)."""

    # ------------------------------------------------------------------
    # Regex patterns (compiled once at import time)
    # ------------------------------------------------------------------
    _GITHUB_RE = re.compile(r"(https?://)?(www\.)?github\.com/[A-Za-z0-9_\-./]+", re.IGNORECASE)
    _LINKEDIN_RE = re.compile(r"(https?://)?(www\.)?linkedin\.com/[A-Za-z0-9_\-./%]+", re.IGNORECASE)

    _EDUCATION_SECTION_RE = re.compile(r"^\s*(education|academic background|qualifications)\s*:?\s*$", re.IGNORECASE)
    _PROJECTS_SECTION_RE = re.compile(r"^\s*(projects?|personal projects?|key projects?)\s*:?\s*$", re.IGNORECASE)
    _CERTS_SECTION_RE = re.compile(r"^\s*(certifications?|licenses?( ?& ?certifications?)?)\s*:?\s*$", re.IGNORECASE)

    # Headers that mark the END of whatever section we're currently reading.
    _SECTION_HEADERS_RE = re.compile(
        r"^\s*(education|experience|work experience|projects?|skills?|"
        r"certifications?|licenses?|summary|objective|references?|"
        r"awards?|publications?|languages?)\s*:?\s*$",
        re.IGNORECASE,
    )

    # Contact-info-style lines (links, email, phone) also terminate a
    # section, since these usually sit in a header/footer block rather
    # than inside Education/Projects/Certifications content.
    _CONTACT_LINE_RE = re.compile(
        r"(github\.com|linkedin\.com|@\S+\.\S+|^\s*(email|phone|contact)\s*:)",
        re.IGNORECASE,
    )

    _DEGREE_KEYWORDS = (
        "bachelor", "master", "b.tech", "btech", "m.tech", "mtech", "mba",
        "b.sc", "bsc", "m.sc", "msc", "phd", "ph.d", "b.e.", "be ", "m.e.",
        "associate degree", "diploma", "bca", "mca",
    )

    _CERT_KEYWORDS = (
        "certified", "certification", "certificate", "aws certified",
        "azure", "google cloud certified", "pmp", "scrum master", "ccna",
        "comptia", "oracle certified", "microsoft certified",
    )

    @staticmethod
    def extract_text(file_path_or_bytes: Union[str, bytes], file_type: str) -> str:
        """Extract text from a document based on its file type.

        Parameters
        ----------
        file_path_or_bytes:
            Either a string file path or a bytes object containing the file.
        file_type:
            The extension or type of the file (e.g., 'pdf', 'docx', 'txt').

        Returns
        -------
        str
            The extracted text.
        """
        file_type = file_type.lower().strip('.')
        text = ""

        try:
            if file_type == 'pdf':
                import fitz  # PyMuPDF
                if isinstance(file_path_or_bytes, bytes):
                    doc = fitz.open(stream=file_path_or_bytes, filetype="pdf")
                else:
                    doc = fitz.open(file_path_or_bytes)
                text = chr(10).join([page.get_text() for page in doc])
                doc.close()

            elif file_type == 'docx':
                import docx
                if isinstance(file_path_or_bytes, bytes):
                    doc = docx.Document(io.BytesIO(file_path_or_bytes))
                else:
                    doc = docx.Document(file_path_or_bytes)
                text = chr(10).join([para.text for para in doc.paragraphs])

            elif file_type == 'txt':
                if isinstance(file_path_or_bytes, bytes):
                    text = file_path_or_bytes.decode('utf-8', errors='ignore')
                else:
                    with open(file_path_or_bytes, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_type} file: {e}")
            raise RuntimeError(f"Failed to extract text from document: {e}")

        return text

    # ------------------------------------------------------------------
    # Structuring
    # ------------------------------------------------------------------

    @staticmethod
    def structure_candidate_data(raw_text: str) -> dict[str, Any]:
        """Structure raw resume text into a standard candidate dictionary.

        Uses lightweight regex / keyword heuristics to extract:
        - Name, Title, Experience, Skills (baseline — unchanged behaviour)
        - GitHub / LinkedIn URLs
        - Education entries
        - Project entries
        - Certification entries

        Parameters
        ----------
        raw_text:
            The raw text extracted from a resume document.

        Returns
        -------
        dict[str, Any]
            Candidate dictionary matching the application's internal schema.
        """
        # Baseline heuristics
        name = "Unknown Candidate"
        title = "Candidate"
        experience_years = 0.0
        skills = []

        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]

        # 1. Attempt to extract Name (assume first line or two has the name)
        if lines:
            name = lines[0][:50]  # Just take the first line up to 50 chars as a guess

        # 2. Attempt to extract Title
        titles_to_look_for = [
            "Software Engineer", "Senior AI Engineer", "Data Scientist",
            "ML Engineer", "Machine Learning Engineer", "Backend Engineer",
            "Data Engineer", "Marketing Manager", "Accountant"
        ]
        for t in titles_to_look_for:
            if re.search(r'\b' + re.escape(t) + r'\b', raw_text, re.IGNORECASE):
                title = t
                break

        # 3. Attempt to extract Experience Years
        exp_match = re.search(r'(\d+(?:\.\d+)?)\+?\s*years?(?:\s*of\s*)?experience', raw_text, re.IGNORECASE)
        if exp_match:
            try:
                experience_years = float(exp_match.group(1))
            except ValueError:
                pass

        # 4. Attempt to extract Skills
        core_skills_to_check = [
            "Python", "PyTorch", "TensorFlow", "Scikit-Learn", "NLP",
            "Computer Vision", "Kubernetes", "Docker", "SQL", "Spark",
            "FastAPI", "MLflow", "LangChain", "CUDA", "Transformers",
            "HuggingFace", "RAG", "LLM", "Machine Learning", "Deep Learning",
            "MLOps", "Ray", "Triton", "Excel", "Java", "C++", "AWS"
        ]

        extracted_skills = []
        for sk in core_skills_to_check:
            if re.search(r'\b' + re.escape(sk) + r'\b', raw_text, re.IGNORECASE):
                extracted_skills.append({
                    "name": sk,
                    "proficiency": "advanced",  # Default guess
                    "duration_months": int(max(12, experience_years * 12 * 0.5))
                })

        if not extracted_skills:
            extracted_skills.append({
                "name": "General IT",
                "proficiency": "intermediate",
                "duration_months": 12
            })

        # 5. GitHub / LinkedIn URLs
        github_url = UniversalParser._first_match(UniversalParser._GITHUB_RE, raw_text)
        linkedin_url = UniversalParser._first_match(UniversalParser._LINKEDIN_RE, raw_text)

        # 6. Education / Projects / Certifications (section-aware, best-effort)
        education = UniversalParser._extract_section_entries(
            lines, UniversalParser._EDUCATION_SECTION_RE, UniversalParser._DEGREE_KEYWORDS
        )
        projects = UniversalParser._extract_project_entries(lines)
        certifications = UniversalParser._extract_section_entries(
            lines, UniversalParser._CERTS_SECTION_RE, UniversalParser._CERT_KEYWORDS
        )

        import uuid
        candidate_id = f"LIVE_{uuid.uuid4().hex[:8].upper()}"

        candidate = {
            "candidate_id": candidate_id,
            "name": name,
            "experience_years": experience_years,
            "profile": {
                "current_title": title,
                "summary": raw_text[:500] + "..." if len(raw_text) > 500 else raw_text,
                "github": github_url or "",
                "linkedin": linkedin_url or "",
            },
            "skills": extracted_skills,
            "experience": [
                {
                    "company": "Live Upload",
                    "title": title,
                    "duration_months": int(experience_years * 12),
                    "description": ""
                }
            ],
            "education": education,
            "projects": projects,
            "certifications": certifications,
            "redrob_signals": {
                "recruiter_response_rate": 1.0,
                "last_active_date": "2026-06-01",
                "interview_completion_rate": 1.0,
            }
        }

        return candidate

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _first_match(pattern: "re.Pattern[str]", text: str) -> str:
        m = pattern.search(text)
        if not m:
            return ""
        url = m.group(0).strip().rstrip(".,;)")
        if not url.lower().startswith("http"):
            url = "https://" + url
        return url

    @staticmethod
    def _extract_section_entries(
        lines: list[str],
        section_header_re: "re.Pattern[str]",
        fallback_keywords: tuple[str, ...],
    ) -> list[str]:
        """Return entries found under a named section header, OR any line
        anywhere in the document matching the fallback keyword list.

        This dual strategy handles both well-structured resumes (clear
        "Education:" / "Certifications:" headers) and loosely-formatted
        ones (keywords scattered inline).
        """
        entries: list[str] = []

        # Strategy A — section header scan
        in_section = False
        for line in lines:
            if section_header_re.match(line):
                in_section = True
                continue
            if in_section:
                if UniversalParser._SECTION_HEADERS_RE.match(line):
                    in_section = False
                    continue
                if UniversalParser._CONTACT_LINE_RE.search(line):
                    continue
                if line and len(line) > 3:
                    entries.append(line[:200])

        # Strategy B — fallback keyword scan (only if section scan found nothing)
        if not entries:
            for line in lines:
                low = line.lower()
                if UniversalParser._CONTACT_LINE_RE.search(line):
                    continue
                if any(kw in low for kw in fallback_keywords):
                    entries.append(line[:200])

        # De-duplicate while preserving order; cap at 10 for UI sanity.
        seen: set[str] = set()
        deduped: list[str] = []
        for e in entries:
            if e not in seen:
                seen.add(e)
                deduped.append(e)
        return deduped[:10]

    @staticmethod
    def _extract_project_entries(lines: list[str]) -> list[str]:
        """Return project entries found under a 'Projects' section header."""
        entries: list[str] = []
        in_section = False
        for line in lines:
            if UniversalParser._PROJECTS_SECTION_RE.match(line):
                in_section = True
                continue
            if in_section:
                if UniversalParser._SECTION_HEADERS_RE.match(line):
                    in_section = False
                    continue
                if UniversalParser._CONTACT_LINE_RE.search(line):
                    continue
                if line and len(line) > 3:
                    entries.append(line[:200])

        seen: set[str] = set()
        deduped: list[str] = []
        for e in entries:
            if e not in seen:
                seen.add(e)
                deduped.append(e)
        return deduped[:10]

# """
# backend/document_parser.py — Multi-Format Extraction Layer
# ===========================================================
# Universal Document Ingestion (Phase 5).

# Supports extracting text and structuring candidate data from live uploaded
# resumes in PDF, DOCX, and TXT formats.
# """

# from __future__ import annotations

# import io
# import logging
# import re
# from typing import Any, Union

# logger = logging.getLogger(__name__)

# class UniversalParser:
#     """Universal parser for resume documents (PDF, DOCX, TXT)."""

#     @staticmethod
#     def extract_text(file_path_or_bytes: Union[str, bytes], file_type: str) -> str:
#         """Extract text from a document based on its file type.

#         Parameters
#         ----------
#         file_path_or_bytes:
#             Either a string file path or a bytes object containing the file.
#         file_type:
#             The extension or type of the file (e.g., 'pdf', 'docx', 'txt').

#         Returns
#         -------
#         str
#             The extracted text.
#         """
#         file_type = file_type.lower().strip('.')
#         text = ""

#         try:
#             if file_type == 'pdf':
#                 import fitz  # PyMuPDF
#                 if isinstance(file_path_or_bytes, bytes):
#                     doc = fitz.open(stream=file_path_or_bytes, filetype="pdf")
#                 else:
#                     doc = fitz.open(file_path_or_bytes)
#                 text = chr(10).join([page.get_text() for page in doc])
#                 doc.close()

#             elif file_type == 'docx':
#                 import docx
#                 if isinstance(file_path_or_bytes, bytes):
#                     doc = docx.Document(io.BytesIO(file_path_or_bytes))
#                 else:
#                     doc = docx.Document(file_path_or_bytes)
#                 text = chr(10).join([para.text for para in doc.paragraphs])

#             elif file_type == 'txt':
#                 if isinstance(file_path_or_bytes, bytes):
#                     text = file_path_or_bytes.decode('utf-8', errors='ignore')
#                 else:
#                     with open(file_path_or_bytes, 'r', encoding='utf-8', errors='ignore') as f:
#                         text = f.read()
#             else:
#                 raise ValueError(f"Unsupported file type: {file_type}")
#         except Exception as e:
#             logger.error(f"Error extracting text from {file_type} file: {e}")
#             raise RuntimeError(f"Failed to extract text from document: {e}")

#         return text

#     @staticmethod
#     def structure_candidate_data(raw_text: str) -> dict[str, Any]:
#         """Structure raw resume text into a standard candidate dictionary.

#         Uses lightweight regex heuristics to extract:
#         - Name
#         - Years of Experience
#         - Skills
#         - Current Title

#         Parameters
#         ----------
#         raw_text:
#             The raw text extracted from a resume document.

#         Returns
#         -------
#         dict[str, Any]
#             Candidate dictionary matching the application's internal schema.
#         """
#         # Baseline heuristics
#         name = "Unknown Candidate"
#         title = "Candidate"
#         experience_years = 0.0
#         skills = []

#         # 1. Attempt to extract Name (assume first line or two has the name)
#         lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
#         if lines:
#             name = lines[0][:50]  # Just take the first line up to 50 chars as a guess

#         # 2. Attempt to extract Title
#         titles_to_look_for = [
#             "Software Engineer", "Senior AI Engineer", "Data Scientist",
#             "ML Engineer", "Machine Learning Engineer", "Backend Engineer",
#             "Data Engineer", "Marketing Manager", "Accountant"
#         ]
#         for t in titles_to_look_for:
#             if re.search(r'\b' + re.escape(t) + r'\b', raw_text, re.IGNORECASE):
#                 title = t
#                 break

#         # 3. Attempt to extract Experience Years
#         # Look for phrases like "5 years of experience" or "5+ years"
#         exp_match = re.search(r'(\d+(?:\.\d+)?)\+?\s*years?(?:\s*of\s*)?experience', raw_text, re.IGNORECASE)
#         if exp_match:
#             try:
#                 experience_years = float(exp_match.group(1))
#             except ValueError:
#                 pass
        
#         # 4. Attempt to extract Skills
#         core_skills_to_check = [
#             "Python", "PyTorch", "TensorFlow", "Scikit-Learn", "NLP",
#             "Computer Vision", "Kubernetes", "Docker", "SQL", "Spark",
#             "FastAPI", "MLflow", "LangChain", "CUDA", "Transformers",
#             "HuggingFace", "RAG", "LLM", "Machine Learning", "Deep Learning",
#             "MLOps", "Ray", "Triton", "Excel", "Java", "C++", "AWS"
#         ]
        
#         extracted_skills = []
#         for sk in core_skills_to_check:
#             if re.search(r'\b' + re.escape(sk) + r'\b', raw_text, re.IGNORECASE):
#                 extracted_skills.append({
#                     "name": sk,
#                     "proficiency": "advanced", # Default guess
#                     "duration_months": int(max(12, experience_years * 12 * 0.5)) # Assign some baseline duration
#                 })
        
#         if not extracted_skills:
#             # Add a fallback so the pipeline doesn't crash on empty skills
#             extracted_skills.append({
#                 "name": "General IT",
#                 "proficiency": "intermediate",
#                 "duration_months": 12
#             })

#         import uuid
#         candidate_id = f"LIVE_{uuid.uuid4().hex[:8].upper()}"

#         # Build the dictionary matching candidate_schema.json
#         candidate = {
#             "candidate_id": candidate_id,
#             "name": name,
#             "experience_years": experience_years,
#             "profile": {
#                 "current_title": title,
#                 "summary": raw_text[:500] + "..." if len(raw_text) > 500 else raw_text,
#                 "github": "",
#                 "linkedin": "",
#             },
#             "skills": extracted_skills,
#             "experience": [
#                 {
#                     "company": "Live Upload",
#                     "title": title,
#                     "duration_months": int(experience_years * 12),
#                     "description": ""
#                 }
#             ],
#             "education": [],
#             "projects": [],
#             "certifications": [],
#             "redrob_signals": {
#                 # Baseline for live uploads to bypass missing signal penalties
#                 "recruiter_response_rate": 1.0,
#                 "last_active_date": "2026-06-01", 
#                 "interview_completion_rate": 1.0,
#             }
#         }

#         return candidate
